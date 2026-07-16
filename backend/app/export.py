"""명세서 초안 내보내기: Markdown / DOCX."""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt

from .schema import SECTIONS


def to_markdown(generated: dict[str, str]) -> str:
    """생성된 섹션들을 하나의 Markdown 문서로."""
    lines = ["# 특허 명세서 (초안)", ""]
    for section in SECTIONS:
        body = (generated.get(section.key) or "").strip()
        if not body:
            continue
        lines.append(f"## {section.title}")
        lines.append("")
        lines.append(body)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def to_docx(generated: dict[str, str]) -> bytes:
    """생성된 섹션들을 DOCX 바이트로."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"  # 한글 폰트
    style.font.size = Pt(10)

    doc.add_heading("특허 명세서 (초안)", level=0)

    for section in SECTIONS:
        body = (generated.get(section.key) or "").strip()
        if not body:
            continue
        doc.add_heading(section.title, level=1)
        for para in body.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
